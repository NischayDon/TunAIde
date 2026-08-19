from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Query, Request, Form
from sqlalchemy.orm import Session
from typing import List, Optional, Dict
import uuid
import os
from io import BytesIO
from docx import Document
from fastapi.responses import StreamingResponse, Response
from pydantic import EmailStr, BaseModel
from fastapi_mail import FastMail, MessageSchema, ConnectionConfig, MessageType
from datetime import datetime, timezone, timedelta

from app.db.base import get_db
from app.db.models import Job, JobStatus, Transcript, User, SupportingDocument
from app.schemas import (
    JobCreate, JobResponse, UploadResponse, TranscriptResponse,
    LedgerEntryUpdate, SupportingDocumentResponse
)
from app.services.storage import storage_service
from app.workers.tasks import process_audio
from app.api.auth import get_current_user
from app.core.config import settings


# Email Configuration — reads from centralized settings (which loads from .env)
conf = ConnectionConfig(
    MAIL_USERNAME=settings.MAIL_USERNAME,
    MAIL_PASSWORD=settings.MAIL_PASSWORD,
    MAIL_FROM=settings.MAIL_FROM,
    MAIL_PORT=settings.MAIL_PORT,
    MAIL_SERVER=settings.MAIL_SERVER,
    MAIL_STARTTLS=settings.MAIL_STARTTLS,
    MAIL_SSL_TLS=settings.MAIL_SSL,
    USE_CREDENTIALS=True,
    VALIDATE_CERTS=True
)

class EmailRequest(BaseModel):
    email: EmailStr

class TranscriptUpdateRequest(BaseModel):
    text_content: str

def generate_docx(job: Job) -> BytesIO:
    document = Document()
    document.add_heading(job.original_filename, 0)
    
    # Logic to populate document
    # Safely access json_metadata, defaulting to empty dict if None
    json_metadata = job.transcript.json_metadata or {}
    segments = json_metadata.get("segments", [])
    
    if segments:
        for seg in segments:
            text = seg.get("text", "")
            document.add_paragraph(text)
    else:
        # Fallback to plain text if no segments available
        text_content = job.transcript.text_content or ""
        for line in text_content.split('\n'):
            if line.strip():
                document.add_paragraph(line)
                
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

router = APIRouter()

@router.post("/upload", response_model=UploadResponse)
def initiate_upload(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    try:
        print(f"Starting upload for file: {file.filename}")
        # Save file to local storage
        saved_filename = storage_service.save_file(file.file, file.filename)
        print(f"File saved successfully at: {saved_filename}")
        
        # Create Job Record with login_date auto-set
        new_job = Job(
            user_id=user.id,
            original_filename=file.filename,
            storage_path=saved_filename,
            status=JobStatus.UPLOADED.value,
            login_date=datetime.now(timezone.utc)
        )
        db.add(new_job)
        
        # Increment lifetime upload counter
        user.total_uploads = (user.total_uploads or 0) + 1
        
        db.commit()
        db.refresh(new_job)
        print(f"Job created successfully: {new_job.id}")

        return UploadResponse(
            upload_url="",
            job_id=new_job.id,
            storage_path=new_job.storage_path
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"CRITICAL UPLOAD FAILURE: {e}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@router.post("/{job_id}/process", response_model=JobResponse)
def start_processing(
    job_id: str, 
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    if job.status != JobStatus.UPLOADED.value and job.status != JobStatus.FAILED.value:
         if job.status in [JobStatus.QUEUED.value, JobStatus.PROCESSING.value, JobStatus.TRANSCRIBING.value]:
             return job

    # Update Status
    job.status = JobStatus.QUEUED.value
    db.commit()
    db.refresh(job)

    # Dispatch to Celery worker via Redis
    process_audio.delay(job_id)

    return job

@router.get("/", response_model=List[JobResponse])
def list_jobs(
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
    skip: int = 0, 
    limit: int = 100,
    status: Optional[str] = Query(None, description="Filter by status"),
    service_type: Optional[str] = Query(None, description="Filter by service type")
):
    query = db.query(Job).filter(Job.user_id == user.id)
    
    if status == "TRASHED":
        query = query.filter(Job.status == JobStatus.TRASHED.value)
    elif status:
        query = query.filter(Job.status == status)
    else:
        # Default: everything NOT trashed
        query = query.filter(Job.status != JobStatus.TRASHED.value)
    
    if service_type:
        query = query.filter(Job.service_type == service_type)

    jobs = query.order_by(Job.created_at.desc()).offset(skip).limit(limit).all()
    return jobs

@router.get("/stats")
def get_lifetime_stats(
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Get lifetime upload/completion stats (survives permanent deletion)."""
    return {
        "total_ever": user.total_uploads or 0,
        "completed_ever": user.total_completed or 0
    }

@router.get("/stats/daily")
def get_daily_stats(
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Get per-day upload/completion counts for the last 7 days (includes all statuses)."""
    now = datetime.now(timezone.utc)
    days = []
    for i in range(6, -1, -1):
        day = now - timedelta(days=i)
        day_start = day.replace(hour=0, minute=0, second=0, microsecond=0)
        day_end = day_start + timedelta(days=1)
        
        uploaded = db.query(Job).filter(
            Job.user_id == user.id,
            Job.created_at >= day_start,
            Job.created_at < day_end
        ).count()
        
        completed = db.query(Job).filter(
            Job.user_id == user.id,
            Job.created_at >= day_start,
            Job.created_at < day_end,
            Job.status == JobStatus.COMPLETED.value
        ).count()
        
        days.append({
            "date": day_start.strftime("%Y-%m-%d"),
            "label": day_start.strftime("%a"),
            "uploaded": uploaded,
            "completed": completed
        })
    
    return days

@router.patch("/{job_id}", response_model=JobResponse)
def update_ledger_entry(
    job_id: str,
    update_data: LedgerEntryUpdate,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Update ledger entry fields for a job."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    # Only update fields that were explicitly set (not None)
    update_dict = update_data.model_dump(exclude_unset=True)
    for field, value in update_dict.items():
        setattr(job, field, value)
    
    db.commit()
    db.refresh(job)
    return job

@router.delete("/trash/all", status_code=204)
def empty_trash(
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """
    Permanently delete all jobs in TRASHED status.
    """
    jobs_to_delete = db.query(Job).filter(
        Job.user_id == user.id, 
        Job.status == JobStatus.TRASHED.value
    ).all()
    
    for job in jobs_to_delete:
        # Delete file from disk
        storage_service.delete_file(job.storage_path)
        # Delete supporting documents from disk
        for doc in job.supporting_documents:
            storage_service.delete_file(doc.storage_path)
        # Delete from DB
        db.delete(job)
        
    db.commit()
    return

@router.get("/{job_id}", response_model=JobResponse)
def get_job(
    job_id: str, 
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return job

@router.get("/{job_id}/transcript", response_model=TranscriptResponse)
def get_transcript(
    job_id: str, 
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    if not job.transcript:
        raise HTTPException(status_code=404, detail="Transcript not ready")
        
    return job.transcript

@router.delete("/{job_id}", response_model=JobResponse)
def delete_job(
    job_id: str, 
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    # Soft delete
    job.status = JobStatus.TRASHED.value
    db.commit()
    db.refresh(job)
    return job

@router.delete("/{job_id}/permanent", status_code=204)
def delete_job_permanent(
    job_id: str, 
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    # Delete audio file from disk
    storage_service.delete_file(job.storage_path)
    # Delete supporting documents from disk
    for doc in job.supporting_documents:
        storage_service.delete_file(doc.storage_path)
    
    # Delete from DB
    db.delete(job)
    db.commit()
    return

@router.post("/{job_id}/restore", response_model=JobResponse)
def restore_job(
    job_id: str, 
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    if job.transcript:
        job.status = JobStatus.COMPLETED.value
    else:
        job.status = JobStatus.FAILED.value
        
    db.commit()
    db.refresh(job)
    return job

@router.post("/{job_id}/email")
async def email_job(
    job_id: str,
    email_req: EmailRequest,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    if not job.transcript:
        raise HTTPException(status_code=400, detail="Transcript not ready")
        
    try:
        # Generate plain text body
        if job.transcript.json_metadata and "segments" in job.transcript.json_metadata:
            segments = job.transcript.json_metadata["segments"]
            body_text = "\n".join([s['text'] for s in segments])
        else:
            # Fallback
            body_text = job.transcript.text_content or "No transcript available."
            
        message = MessageSchema(
            subject=f"Transcript: {job.original_filename}",
            recipients=[email_req.email],
            body=body_text,
            subtype=MessageType.plain
        )
        
        fm = FastMail(conf)
        await fm.send_message(message)
        
        return {"message": "Email sent successfully"}
    except Exception as e:
        print(f"Email failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to send email: {str(e)}")

@router.get("/{job_id}/download")
def download_job(
    job_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    if not job.transcript:
        raise HTTPException(status_code=400, detail="Transcript not ready")
        
    buffer = generate_docx(job)
    filename = f"{job.original_filename}.docx"
    
    return StreamingResponse(
        buffer, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.get("/{job_id}/audio")
def stream_audio(
    job_id: str,
    request: Request,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Stream the original uploaded audio file for browser playback."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    # Download from cloud storage to temp file
    try:
        file_path = storage_service.download_to_temp(job.storage_path)
    except Exception as e:
        print(f"Failed to download audio for streaming: {e}")
        raise HTTPException(status_code=500, detail="Failed to load audio file")
    
    # Determine MIME type from extension
    ext = job.original_filename.rsplit(".", 1)[-1].lower() if "." in job.original_filename else "bin"
    audio_mime_map = {
        "mp3": "audio/mpeg",
        "wav": "audio/wav",
        "ogg": "audio/ogg",
        "flac": "audio/flac",
        "aac": "audio/aac",
        "m4a": "audio/mp4",
        "wma": "audio/x-ms-wma",
        "webm": "audio/webm",
        "mp4": "video/mp4",
        "mkv": "video/x-matroska",
        "avi": "video/x-msvideo",
        "mov": "video/quicktime",
    }
    media_type = audio_mime_map.get(ext, "application/octet-stream")
    
    file_size = os.path.getsize(file_path)
    
    # Support Range header for seeking
    range_header = request.headers.get("range")
    if range_header:
        # Parse "bytes=start-end"
        range_spec = range_header.replace("bytes=", "")
        parts = range_spec.split("-")
        start = int(parts[0]) if parts[0] else 0
        end = int(parts[1]) if parts[1] else file_size - 1
        
        if start >= file_size:
            raise HTTPException(status_code=416, detail="Range not satisfiable")
        
        end = min(end, file_size - 1)
        content_length = end - start + 1
        
        def iter_range():
            with open(file_path, "rb") as f:
                f.seek(start)
                remaining = content_length
                while remaining > 0:
                    chunk_size = min(8192, remaining)
                    data = f.read(chunk_size)
                    if not data:
                        break
                    remaining -= len(data)
                    yield data
        
        return Response(
            content=b"".join(iter_range()),
            status_code=206,
            media_type=media_type,
            headers={
                "Content-Range": f"bytes {start}-{end}/{file_size}",
                "Accept-Ranges": "bytes",
                "Content-Length": str(content_length),
            }
        )
    else:
        # Full file response
        def iterfile():
            with open(file_path, "rb") as f:
                yield from f
        
        return StreamingResponse(
            iterfile(),
            media_type=media_type,
            headers={
                "Accept-Ranges": "bytes",
                "Content-Length": str(file_size),
            }
        )

@router.put("/{job_id}/transcript", response_model=TranscriptResponse)
def update_transcript(
    job_id: str,
    update_data: TranscriptUpdateRequest,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Update transcript text content (from inline editor)."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    if not job.transcript:
        raise HTTPException(status_code=404, detail="Transcript not found")
    
    # Update text content
    job.transcript.text_content = update_data.text_content
    
    # Rebuild segments from updated text if segments existed
    if job.transcript.json_metadata and "segments" in job.transcript.json_metadata:
        # Replace all segment text with the full updated text as a single segment
        job.transcript.json_metadata = {
            **job.transcript.json_metadata,
            "segments": [{"text": update_data.text_content, "start": "00:00:00", "end": "--:--:--"}]
        }
    
    db.commit()
    db.refresh(job.transcript)
    return job.transcript


# =====================================================
# Supporting Documents Endpoints
# =====================================================

@router.post("/{job_id}/documents", response_model=SupportingDocumentResponse)
def upload_supporting_document(
    job_id: str,
    file: UploadFile = File(...),
    description: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Upload a supporting document for a ledger entry."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    try:
        # Save to storage
        saved_filename = storage_service.save_file(file.file, file.filename)
        
        # Get file size
        file.file.seek(0, 2)  # Seek to end
        file_size = file.file.tell()
        file.file.seek(0)
        
        doc = SupportingDocument(
            job_id=job_id,
            original_filename=file.filename,
            storage_path=saved_filename,
            file_size_bytes=file_size,
            description=description
        )
        db.add(doc)
    except Exception as e:
        print(f"Failed to upload supporting document {file.filename}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to upload {file.filename}: {str(e)}")
    
    db.commit()
    db.refresh(doc)
    
    return doc

@router.get("/{job_id}/documents", response_model=List[SupportingDocumentResponse])
def list_supporting_documents(
    job_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """List all supporting documents for a ledger entry."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    return job.supporting_documents

@router.get("/{job_id}/documents/{doc_id}/download")
def download_supporting_document(
    job_id: str,
    doc_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Download a specific supporting document."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    doc = db.query(SupportingDocument).filter(
        SupportingDocument.id == doc_id,
        SupportingDocument.job_id == job_id
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    file_path = storage_service.download_to_temp(doc.storage_path)
    
    # Determine MIME type from extension
    ext = doc.original_filename.rsplit(".", 1)[-1].lower() if "." in doc.original_filename else "bin"
    mime_map = {
        "pdf": "application/pdf",
        "doc": "application/msword",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "jpg": "image/jpeg", "jpeg": "image/jpeg",
        "png": "image/png",
        "txt": "text/plain",
    }
    media_type = mime_map.get(ext, "application/octet-stream")
    
    def iterfile():
        with open(file_path, "rb") as f:
            yield from f
    
    return StreamingResponse(
        iterfile(),
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={doc.original_filename}"}
    )

@router.delete("/{job_id}/documents/{doc_id}", status_code=204)
def delete_supporting_document(
    job_id: str,
    doc_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """Delete a specific supporting document."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    doc = db.query(SupportingDocument).filter(
        SupportingDocument.id == doc_id,
        SupportingDocument.job_id == job_id
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Delete file from storage
    storage_service.delete_file(doc.storage_path)
    
    # Delete from DB
    db.delete(doc)
    db.commit()
    return

